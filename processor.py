import re

from docx_notes import preserve_notes
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

PROFILE_SETTINGS = {
    "novel": {
        "font_name": "Garamond",
        "font_size": 12,
        "line_spacing": 1.15,
        "first_line_indent_cm": 1,
        "normalize_slash_spacing": False,
        "protect_patterns": [
            r"^[A-Z][a-z]+:$",
        ],
        "merge_after_punctuation": False,
    },
    "academic": {
        "font_name": "Arial",
        "font_size": 11,
        "line_spacing": 1.15,
        "first_line_indent_cm": 1,
        "normalize_slash_spacing": True,
        "protect_patterns": [],
        "merge_after_punctuation": False,
    },
    "legal": {
        "font_name": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.0,
        "first_line_indent_cm": 0,
        "normalize_slash_spacing": False,
        "protect_patterns": [
            r"^(Article|Section|Clause|Paragraph)\s+\d+",
            r"^§+\s*\d+",
            r"^\(\d+\)",
            r"^\d+\.\d+",
        ],
        "merge_after_punctuation": False,
    },
}

QUOTE_STYLES = {
    "english-double": ('"', '"'),
    "english-single": ("'", "'"),
    "serbian": ("„", "”"),
    "german": ("„", "“"),
}

LIGATURE_MAP = str.maketrans(
    {
        "ﬁ": "fi",
        "ﬂ": "fl",
        "ﬀ": "ff",
        "ﬃ": "ffi",
        "ﬄ": "ffl",
        "ﬅ": "ft",
        "ﬆ": "st",
    }
)

DOUBLE_QUOTES_PATTERN = re.compile(r"[“”„‟«»]")
SINGLE_QUOTES_PATTERN = re.compile(r"[‘’‚‛‹›]")
LETTER_PATTERN = r"[^\W\d_]"


class CleaningCancelled(Exception):
    pass


PROTECTED_RUN_XPATHS = (
    ".//w:footnoteReference",
    ".//w:endnoteReference",
    ".//w:commentReference",
    ".//w:fldChar",
    ".//w:instrText",
)


def iter_table_paragraph_collections(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.paragraphs
                yield from iter_table_paragraph_collections(cell.tables)


def iter_paragraph_collections(doc):
    yield doc.paragraphs
    yield from iter_table_paragraph_collections(doc.tables)


def iter_all_paragraphs(doc):
    for paragraphs in iter_paragraph_collections(doc):
        for paragraph in paragraphs:
            yield paragraph


def collect_all_paragraphs(doc):
    return list(iter_all_paragraphs(doc))


def run_has_protected_markup(run):
    return any(run._element.xpath(xpath) for xpath in PROTECTED_RUN_XPATHS)


def paragraph_has_protected_markup(paragraph):
    return any(run_has_protected_markup(run) for run in paragraph.runs)


def ensure_not_cancelled(should_cancel):
    if should_cancel is not None and should_cancel():
        raise CleaningCancelled("Operation cancelled.")


def normalize_false_number_spacing(txt):
    while True:
        txt, count = re.subn(r"(?<=\d)\s+(?=\d{3}\b)", "", txt)
        if count == 0:
            return txt


def normalize_quotes(txt):
    txt = txt.replace("``", '"')
    txt = DOUBLE_QUOTES_PATTERN.sub('"', txt)
    txt = SINGLE_QUOTES_PATTERN.sub("'", txt)
    txt = re.sub(r"(?<!\w)''(?=\w)", '"', txt)
    txt = re.sub(r'(?<=\w)""(?=[\s,.!?;:)]|$)', '"', txt)
    txt = re.sub(r"(?<=\w)''(?=[\s,.!?;:)]|$)", '"', txt)
    return txt


def normalize_duplicate_punctuation(txt):
    txt = re.sub(r"([,;:])(?:\s*\1)+", r"\1", txt)
    txt = re.sub(r"([!?])(?:\s*\1)+", r"\1", txt)
    txt = re.sub(r"\.(?:\s*\.){1,}", ".", txt)
    txt = re.sub(r":\s*,", ":", txt)
    txt = re.sub(r";\s*,", ";", txt)
    return txt


def normalize_broken_word_hyphenation(txt):
    return re.sub(rf"(?<={LETTER_PATTERN})\s*-\s+(?={LETTER_PATTERN})", "", txt)


def normalize_special_spacing(txt, normalize_slash_spacing=True):
    if normalize_slash_spacing:
        txt = re.sub(r"(?<=\w)\s*/\s*(?=\w)", "/", txt)
    txt = re.sub(r"(?<=\d)\s+%(?=\W|$)", "%", txt)
    return txt


def normalize_double_quote_spacing(txt):
    result = []
    inside_quotes = False
    i = 0
    while i < len(txt):
        char = txt[i]
        if char == '"':
            if inside_quotes:
                while result and result[-1] == " ":
                    result.pop()
                result.append(char)
                inside_quotes = False
                i += 1
                continue

            result.append(char)
            inside_quotes = True
            i += 1
            while i < len(txt) and txt[i] == " ":
                i += 1
            continue

        result.append(char)
        i += 1

    return "".join(result)


def normalize_ocr_closing_quote_11(txt):
    # Common OCR artifact: closing quote recognized as 11 after an opening quote.
    return re.sub(r'("([^"\r\n]{1,120}))11(?=(?:\s|$|[,.!?;:)\]}]))', r'\1"', txt)


def normalize_quote_boundaries(txt):
    # Add a missing space before an opening quote glued to the previous word.
    txt = re.sub(r'(?<=[^\s"(\[{])"(?=[A-Za-zА-Яа-яČĆŽŠĐčćžšđ])', ' "', txt)
    return txt


def apply_quote_style_to_text(txt, quote_language):
    opening_quote, closing_quote = QUOTE_STYLES.get(quote_language, QUOTE_STYLES["english-double"])
    result = []
    inside_quotes = False
    for char in txt:
        if char == '"':
            result.append(closing_quote if inside_quotes else opening_quote)
            inside_quotes = not inside_quotes
        else:
            result.append(char)
    return "".join(result)


def apply_quote_style_to_segments(segments, quote_language):
    opening_quote, closing_quote = QUOTE_STYLES.get(quote_language, QUOTE_STYLES["english-double"])
    inside_quotes = False
    normalized_segments = []

    for segment in segments:
        result = []
        for char in segment:
            if char == '"':
                result.append(closing_quote if inside_quotes else opening_quote)
                inside_quotes = not inside_quotes
            else:
                result.append(char)
        normalized_segments.append("".join(result))

    return normalized_segments


def clean_spacing_in_run(run, profile_name, quote_style='"'):
    if run_has_protected_markup(run):
        return
    run.text = normalize_run_text(run.text, profile_name, quote_style)


def normalize_run_text(text, profile_name, quote_style='"'):
    settings = get_profile_settings(profile_name)
    txt = text.translate(LIGATURE_MAP)
    txt = txt.replace("\t", " ")
    txt = normalize_quotes(txt)
    txt = normalize_ocr_closing_quote_11(txt)
    txt = normalize_quote_boundaries(txt)
    txt = normalize_broken_word_hyphenation(txt)
    txt = re.sub(r"(?<=\S)\s*[–—]\s*(?=\S)", " - ", txt)
    txt = txt.replace("–", "-").replace("—", "-")
    txt = normalize_false_number_spacing(txt)
    txt = normalize_special_spacing(txt, settings["normalize_slash_spacing"])
    txt = re.sub(r" {2,}", " ", txt)
    txt = re.sub(r"\s+([,.;:!?])", r"\1", txt)
    txt = re.sub(r"([(\[{])\s+", r"\1", txt)
    txt = re.sub(r"\s+([)\]}])", r"\1", txt)
    txt = normalize_double_quote_spacing(txt)
    txt = normalize_duplicate_punctuation(txt)
    if quote_style == '"':
        txt = txt.replace("''", '"')
    return txt


def get_profile_settings(profile_name):
    return PROFILE_SETTINGS.get(profile_name, PROFILE_SETTINGS["academic"])


def report_progress(progress_callback, percent, message):
    if progress_callback is not None:
        progress_callback(max(0, min(100, int(percent))), message)


def make_stage_reporter(progress_callback, start_percent, end_percent, label):
    def stage_report(done, total):
        if total <= 0:
            percent = end_percent
        else:
            fraction = done / total
            percent = start_percent + (end_percent - start_percent) * fraction
        report_progress(progress_callback, percent, label)

    return stage_report


def for_each_paragraph(paragraphs, callback, progress=None, should_cancel=None):
    total = len(paragraphs)
    if total == 0:
        if progress is not None:
            progress(1, 1)
        return

    for index, paragraph in enumerate(paragraphs, start=1):
        ensure_not_cancelled(should_cancel)
        callback(paragraph)
        if progress is not None and (index == total or index % 10 == 0):
            progress(index, total)


def delete_empty_paragraphs(doc, log, progress=None, should_cancel=None):
    removed = 0
    paragraphs_to_remove = []
    paragraphs = collect_all_paragraphs(doc)
    total = len(paragraphs)

    for index, paragraph in enumerate(paragraphs, start=1):
        ensure_not_cancelled(should_cancel)
        if paragraph.text.strip() == "":
            paragraphs_to_remove.append(paragraph)
        if progress is not None and (index == total or index % 10 == 0):
            progress(index, total)

    for paragraph in paragraphs_to_remove:
        paragraph._element.getparent().remove(paragraph._element)
        removed += 1

    log(f"  - Removed {removed} blank paragraphs")
    return removed


def remove_breaks(doc, log, progress=None, should_cancel=None):
    page_breaks = 0
    paragraphs = collect_all_paragraphs(doc)
    total = len(paragraphs)

    for index, paragraph in enumerate(paragraphs, start=1):
        ensure_not_cancelled(should_cancel)
        for run in paragraph.runs:
            for br in run._element.findall(".//w:br", run._element.nsmap):
                if br.get(qn("w:type")) in ("page", "column", "textWrapping"):
                    br.getparent().remove(br)
                    page_breaks += 1
        if progress is not None and (index == total or index % 10 == 0):
            progress(index, total)

    for section in doc.sections[1:]:
        section.start_type = WD_SECTION_START.CONTINUOUS
    log(f"  - Deleted {page_breaks} manual breaks; sections set to continuous")
    return page_breaks


def reset_indents(doc, log, profile_name, progress=None, should_cancel=None):
    settings = get_profile_settings(profile_name)
    paragraphs = collect_all_paragraphs(doc)

    def apply_indent(paragraph):
        fmt = paragraph.paragraph_format
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        is_heading = style_name.startswith(("heading", "title"))
        fmt.left_indent = Inches(0)
        fmt.right_indent = Inches(0)
        fmt.first_line_indent = None if is_heading else Cm(settings["first_line_indent_cm"])

    for_each_paragraph(paragraphs, apply_indent, progress, should_cancel)
    log(f"  - Indents reset (left/right 0 cm; body first-line {settings['first_line_indent_cm']} cm)")
    return len(paragraphs)


def unify_body_text(doc, log, profile_name, progress=None, should_cancel=None):
    settings = get_profile_settings(profile_name)
    paragraphs = collect_all_paragraphs(doc)
    body_count = 0

    def apply_style(paragraph):
        nonlocal body_count
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if style_name.startswith(("heading", "title")):
            return
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = settings["line_spacing"]
        for run in paragraph.runs:
            if run_has_protected_markup(run):
                continue
            run.font.name = settings["font_name"]
            run.font.size = Pt(settings["font_size"])
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn("w:ascii"), settings["font_name"])
                run._element.rPr.rFonts.set(qn("w:hAnsi"), settings["font_name"])
                run._element.rPr.rFonts.set(qn("w:cs"), settings["font_name"])
                run._element.rPr.rFonts.set(qn("w:eastAsia"), settings["font_name"])
        body_count += 1

    for_each_paragraph(paragraphs, apply_style, progress, should_cancel)
    log(
        "  - Unified body text on "
        f"{body_count} paragraphs ({settings['font_name']} {settings['font_size']}, "
        f"spacing {settings['line_spacing']})"
    )
    return body_count


def uniform_quotes(doc, log, quote_language, progress=None, should_cancel=None):
    paragraphs = collect_all_paragraphs(doc)
    changed_runs = 0

    def apply_quotes(paragraph):
        nonlocal changed_runs
        if not paragraph.runs:
            return
        quote_segments = []
        index_map = []
        for idx, run in enumerate(paragraph.runs):
            if run_has_protected_markup(run):
                continue
            quote_segments.append(run.text)
            index_map.append(idx)
        normalized_segments = apply_quote_style_to_segments(quote_segments, quote_language)
        for idx, normalized_text in zip(index_map, normalized_segments):
            if paragraph.runs[idx].text != normalized_text:
                changed_runs += 1
            paragraph.runs[idx].text = normalized_text

    for_each_paragraph(paragraphs, apply_quotes, progress, should_cancel)
    log(f"  - Uniformed quotes using {quote_language} style")
    return changed_runs


def clone_run(src_run, dest_paragraph, prepend=""):
    new_run = dest_paragraph.add_run(prepend + src_run.text)
    new_run.bold = src_run.bold
    new_run.italic = src_run.italic
    new_run.underline = src_run.underline
    if src_run.font.name:
        new_run.font.name = src_run.font.name
    if src_run.font.size:
        new_run.font.size = src_run.font.size
    if src_run.font.color.rgb:
        new_run.font.color.rgb = src_run.font.color.rgb
    return new_run


def should_skip_merge(current_paragraph, next_paragraph, profile_name):
    settings = get_profile_settings(profile_name)
    if current_paragraph.text.strip() == "" or next_paragraph.text.strip() == "":
        return True

    if paragraph_has_protected_markup(current_paragraph) or paragraph_has_protected_markup(next_paragraph):
        return True

    current_style = current_paragraph.style.name.lower() if current_paragraph.style else ""
    next_style = next_paragraph.style.name.lower() if next_paragraph.style else ""
    if current_style.startswith(("heading", "title")) or next_style.startswith(("heading", "title")):
        return True

    if current_paragraph._element.xpath("./w:pPr/w:numPr") or next_paragraph._element.xpath("./w:pPr/w:numPr"):
        return True

    current_text = current_paragraph.text.strip()
    next_text = next_paragraph.text.strip()
    if not current_text or not next_text:
        return True

    if len(current_text) <= 2 or len(next_text) <= 2:
        return True

    if re.match(r"^[-*•]\s", current_text) or re.match(r"^[-*•]\s", next_text):
        return True

    if re.match(r"^\(?\d+[.)]\s", current_text) or re.match(r"^\(?\d+[.)]\s", next_text):
        return True

    if re.match(r"^[IVXLCDM]+\.\s", current_text) or re.match(r"^[IVXLCDM]+\.\s", next_text):
        return True

    if re.match(r"^[A-Z0-9][A-Z0-9\s.:/-]*$", next_text):
        return True

    for pattern in settings["protect_patterns"]:
        if re.match(pattern, current_text) or re.match(pattern, next_text):
            return True

    return False


def should_merge_paragraphs(current_paragraph, next_paragraph, profile_name):
    settings = get_profile_settings(profile_name)
    current_text = current_paragraph.text.rstrip()
    next_text = next_paragraph.text.lstrip()
    if not current_text or not next_text:
        return False, False

    last_char = current_text[-1]
    if last_char == "-":
        return True, True

    next_core = next_text.lstrip("\"' ([{-")
    if not next_core:
        return False, False

    if last_char in ".!?;:)]}\"'":
        if not settings["merge_after_punctuation"]:
            return False, False

    if next_core[0].islower():
        return True, False

    return False, False


def fix_broken_sentences_in_collection(paragraphs, profile_name, progress=None, should_cancel=None):
    merges = 0
    total = len(paragraphs)
    i = 0
    while i < len(paragraphs) - 1:
        ensure_not_cancelled(should_cancel)
        current_paragraph = paragraphs[i]
        next_paragraph = paragraphs[i + 1]

        if should_skip_merge(current_paragraph, next_paragraph, profile_name):
            i += 1
            if progress is not None and (i >= len(paragraphs) - 1 or i % 10 == 0):
                progress(min(i, total), max(total, 1))
            continue

        join, strip_hyphen = should_merge_paragraphs(current_paragraph, next_paragraph, profile_name)

        if join:
            if strip_hyphen and current_paragraph.runs:
                current_paragraph.runs[-1].text = current_paragraph.runs[-1].text.rstrip("-")

            if not strip_hyphen and not current_paragraph.text.endswith(" "):
                current_paragraph.add_run(" ")

            for run in next_paragraph.runs:
                clone_run(run, current_paragraph)

            next_paragraph._element.getparent().remove(next_paragraph._element)
            paragraphs.pop(i + 1)
            merges += 1
        else:
            i += 1

        if progress is not None and (i >= len(paragraphs) - 1 or i % 10 == 0):
            progress(min(i + 1, total), max(total, 1))

    if progress is not None:
        progress(max(total, 1), max(total, 1))
    return merges


def fix_broken_sentences(doc, log, profile_name, progress=None, should_cancel=None):
    merges = 0
    collections = list(iter_paragraph_collections(doc))
    total = len(collections)
    if total == 0:
        log("  - Merged 0 broken sentence pairs")
        return

    for index, paragraph_collection in enumerate(collections, start=1):
        ensure_not_cancelled(should_cancel)
        local_progress = None
        if progress is not None:
            start_fraction = (index - 1) / total
            end_fraction = index / total

            def local_progress(done, inner_total, start_fraction=start_fraction, end_fraction=end_fraction):
                if inner_total <= 0:
                    progress(index, total)
                    return
                nested_fraction = done / inner_total
                overall_done = start_fraction + (end_fraction - start_fraction) * nested_fraction
                progress(overall_done, 1)

        merges += fix_broken_sentences_in_collection(
            list(paragraph_collection), profile_name, local_progress, should_cancel
        )
        if progress is not None:
            progress(index, total)

    log(f"  - Merged {merges} broken sentence pairs")
    return merges


def process_docx(
    src,
    dst,
    do_spacing,
    do_blanks,
    do_breaks,
    do_indents,
    do_unify,
    do_sentfix,
    do_quote_uniform,
    quote_language,
    profile_name,
    log,
    progress_callback=None,
    should_cancel=None,
    audit_log=None,
):
    report_progress(progress_callback, 0, "Opening file")
    ensure_not_cancelled(should_cancel)
    log(f"Opening file: {src}")
    doc = Document(src)

    enabled_stages = []
    if do_spacing:
        enabled_stages.append(
            ("Clean spacing", lambda reporter: _run_spacing(doc, profile_name, reporter, should_cancel, audit_log))
        )
    if do_blanks:
        enabled_stages.append(
            ("Delete blank rows", lambda reporter: _run_delete_blanks(doc, log, reporter, should_cancel, audit_log))
        )
    if do_breaks:
        enabled_stages.append(
            ("Remove breaks", lambda reporter: _run_remove_breaks(doc, log, reporter, should_cancel, audit_log))
        )
    if do_indents:
        enabled_stages.append(
            ("Reset indents", lambda reporter: _run_reset_indents(doc, log, profile_name, reporter, should_cancel, audit_log))
        )
    if do_unify:
        enabled_stages.append(
            ("Unify body text", lambda reporter: _run_unify_body_text(doc, log, profile_name, reporter, should_cancel, audit_log))
        )
    if do_sentfix:
        enabled_stages.append(
            ("Fix broken sentences", lambda reporter: _run_fix_broken_sentences(doc, log, profile_name, reporter, should_cancel, audit_log))
        )
    if do_quote_uniform:
        enabled_stages.append(
            ("Uniform quotes", lambda reporter: _run_uniform_quotes(doc, log, quote_language, reporter, should_cancel, audit_log))
        )

    total_stages = len(enabled_stages)
    start_base = 5
    end_base = 92
    stage_span = (end_base - start_base) / total_stages if total_stages else 0

    for index, (label, action) in enumerate(enabled_stages):
        ensure_not_cancelled(should_cancel)
        log(f"Stage: {label}")
        stage_start = start_base + index * stage_span
        stage_end = stage_start + stage_span
        reporter = make_stage_reporter(progress_callback, stage_start, stage_end, label)
        reporter(0, 1)
        action(reporter)
        reporter(1, 1)

    ensure_not_cancelled(should_cancel)
    report_progress(progress_callback, 95, "Saving file")
    log("Saving...")
    doc.save(dst)
    preserved_parts = preserve_notes(src, dst)
    if audit_log is not None and preserved_parts:
        audit_log.add_note(f"Preserved package parts: {', '.join(sorted(preserved_parts))}")
    log(f"Done. Saved to: {dst}")
    report_progress(progress_callback, 100, "Finished")


def _run_spacing(doc, profile_name, progress, should_cancel, audit_log):
    paragraphs = collect_all_paragraphs(doc)

    def clean_paragraph(paragraph):
        for run_index, run in enumerate(paragraph.runs, start=1):
            if run_has_protected_markup(run):
                continue
            before = run.text
            after = normalize_run_text(before, profile_name)
            if before != after and audit_log is not None:
                context = paragraph.text.strip()[:120]
                audit_log.record_change("text_normalization", before, after, context=f"run {run_index}: {context}")
            run.text = after

    for_each_paragraph(paragraphs, clean_paragraph, progress, should_cancel)


def _run_delete_blanks(doc, log, progress, should_cancel, audit_log):
    removed = delete_empty_paragraphs(doc, log, progress, should_cancel)
    if audit_log is not None:
        audit_log.bump("blank_paragraphs_removed", removed)


def _run_remove_breaks(doc, log, progress, should_cancel, audit_log):
    removed = remove_breaks(doc, log, progress, should_cancel)
    if audit_log is not None:
        audit_log.bump("breaks_removed", removed)


def _run_reset_indents(doc, log, profile_name, progress, should_cancel, audit_log):
    count = reset_indents(doc, log, profile_name, progress, should_cancel)
    if audit_log is not None:
        audit_log.bump("paragraphs_with_indents_reset", count)


def _run_unify_body_text(doc, log, profile_name, progress, should_cancel, audit_log):
    count = unify_body_text(doc, log, profile_name, progress, should_cancel)
    if audit_log is not None:
        audit_log.bump("paragraphs_unified", count)


def _run_fix_broken_sentences(doc, log, profile_name, progress, should_cancel, audit_log):
    if audit_log is None:
        fix_broken_sentences(doc, log, profile_name, progress, should_cancel)
        return

    paragraphs = list(iter_paragraph_collections(doc))
    # lightweight snapshot before running the merge stage
    before_snapshot = [
        [paragraph.text for paragraph in collection]
        for collection in paragraphs
    ]
    merges = fix_broken_sentences(doc, log, profile_name, progress, should_cancel)
    after_snapshot = [
        [paragraph.text for paragraph in collection]
        for collection in list(iter_paragraph_collections(doc))
    ]
    audit_log.bump("merged_paragraph_pairs", merges)
    for before_collection, after_collection in zip(before_snapshot, after_snapshot):
        if len(before_collection) <= len(after_collection):
            continue
        for index in range(len(before_collection) - 1):
            if index >= len(after_collection):
                break
            before_first = before_collection[index].strip()
            before_second = before_collection[index + 1].strip()
            combined_before = f"{before_first} || {before_second}"
            after_value = after_collection[index].strip()
            if after_value != before_first and after_value != before_second:
                audit_log.record_change("paragraph_merge", combined_before, after_value)


def _run_uniform_quotes(doc, log, quote_language, progress, should_cancel, audit_log):
    if audit_log is None:
        uniform_quotes(doc, log, quote_language, progress, should_cancel)
        return

    paragraphs = collect_all_paragraphs(doc)

    def apply_quotes(paragraph):
        if not paragraph.runs:
            return
        quote_segments = []
        index_map = []
        for idx, run in enumerate(paragraph.runs):
            if run_has_protected_markup(run):
                continue
            quote_segments.append(run.text)
            index_map.append(idx)
        normalized_segments = apply_quote_style_to_segments(quote_segments, quote_language)
        for idx, normalized_text in zip(index_map, normalized_segments):
            run = paragraph.runs[idx]
            before = run.text
            if before != normalized_text:
                context = paragraph.text.strip()[:120]
                audit_log.record_change("quote_uniformization", before, normalized_text, context=context)
                run.text = normalized_text

    for_each_paragraph(paragraphs, apply_quotes, progress, should_cancel)
    log(f"  - Uniformed quotes using {quote_language} style")
